import os
from docx import Document
from copy import deepcopy
from data_engine import getTemptureForOL, getPressureForOL, getFlowForOL, getLevelForOL, get_spec, get_io, get_tsp, \
    get_kj
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from paths import SETTINGS


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def cell_style(cell, size):
    paragra_ph = cell.paragraphs[0]
    paragra_ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    font = paragra_ph.runs[0].font
    font.name = 'Arial'
    font.size = Pt(size)


# def unload(parametr, df_table_list_positions, df_ol_table):
#     global savePath
#     if parametr == 'Температура':
#         document = Document(f'Выгрузка/Шаблоны/ОЛ/Температура.docx')
#         savePath = '193-РП-АТХ1.ОЛ1'
#     elif parametr == 'Давление':
#         document = Document(f'Выгрузка/Шаблоны/ОЛ/Давление.docx')
#         savePath = '193-РП-АТХ1.ОЛ2'
#     elif parametr == 'Расход':
#         document = Document(f'Выгрузка/Шаблоны/ОЛ/Расход.docx')
#         savePath = '193-РП-АТХ1.ОЛ3'
#     elif parametr == 'Уровень':
#         document = Document(f'Выгрузка/Шаблоны/ОЛ/Уровень.docx')
#         savePath = '193-РП-АТХ1.ОЛ4'
#     else:
#         return
#     table_list_positions = document.tables[-2]
#
#     for index, row in df_table_list_positions.iterrows():
#         tb_row = table_list_positions.add_row()
#         cells = tb_row.cells
#
#         for _ in range(0, len(cells)):
#             cell = cells[_]
#             cell.text = str(row.iloc[_])
#             cell_style(cell, 10)
#
#     ol_table = document.tables[-1]
#     last_humans_paragraph = document.paragraphs[-2]
#
#     for index, row in df_ol_table.iterrows():
#         document.add_page_break()
#         # paragraph = document.add_paragraph()
#         paragraph = document.paragraphs[-1]
#         paragraph._p.addnext(deepcopy(ol_table._tbl))
#         cur_table = document.tables[-1]
#         for item in templates[parametr].keys():
#             n_row, n_col = templates[parametr][item]
#             cell = cur_table.cell(n_row, n_col)
#             cell.text = str(row[item])
#             if n_col > 4:
#                 cell_style(cell, 10)
#     ol_table._element.getparent().remove(ol_table._element)
#     delete_paragraph(last_humans_paragraph)
#
#     document_env = Document(f'Шаблоны/ОЛ/Среды.docx')
#     table_env = document_env.tables[-1]
#
#     document.add_page_break()
#     paragraph = document.add_paragraph('Приложение 1')
#
#     paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#     paragraph._p.addnext(deepcopy(table_env._tbl))
#
#
#     document.save(f'Выгрузка/Опросные листы/{savePath}.docx')

def unload_ol(template_path, template, save_path, save_name, data_func):
    df_table_list_positions, df_ol_table = data_func()
    full_save_path = os.path.join(save_path, save_name + '.docx')

    # if not (os.path.exists(template_path)):
    #     print(f'Директории {template_path} для шаблона "{template_name}" не существует')
    #     return

    document = Document(template_path)
    table_list_positions = document.tables[-2]

    for index, row in df_table_list_positions.iterrows():
        tb_row = table_list_positions.add_row()
        cells = tb_row.cells

        for _ in range(0, len(cells)):
            cell = cells[_]
            cell.text = str(row.iloc[_])
            cell_style(cell, 10)

    ol_table = document.tables[-1]
    last_humans_paragraph = document.paragraphs[-2]

    for index, row in df_ol_table.iterrows():
        document.add_page_break()
        # paragraph = document.add_paragraph()
        paragraph = document.paragraphs[-1]
        paragraph._p.addnext(deepcopy(ol_table._tbl))
        cur_table = document.tables[-1]
        for item in template.keys():
            n_row, n_col = template[item]
            cell = cur_table.cell(n_row, n_col)
            cell.text = str(row[item])
            if n_col > 4:
                cell_style(cell, 10)
    ol_table._element.getparent().remove(ol_table._element)
    delete_paragraph(last_humans_paragraph)

    document_env = Document(SETTINGS["file_dir_environments_descriptions"])
    table_env = document_env.tables[-1]

    document.add_page_break()
    paragraph = document.add_paragraph('Приложение 1')

    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    paragraph._p.addnext(deepcopy(table_env._tbl))

    document.save(full_save_path)

# unload_ol(
#                 template_path=SETTINGS["file_dir_flow_template"],
#                 template=SETTINGS["template_flow"],
#                 save_path=SETTINGS["dir_ol_save_directory"],
#                 save_name=SETTINGS["file_name_flow_ol"],
#                 data_func=getFlowForOL)
# t_ol = getTemptureForOL()
# print('HI')
#
#
# # unload_ol(template_path, template_name, save_path, save_name, df_table_list_positions, df_ol_table)
# unload_ol(template_path=SETTINGS["file_dir_temperature_template"],
#           template=SETTINGS["template_temperature"],
#           save_path=SETTINGS["dir_ol_save_directory"],
#           save_name=SETTINGS["file_name_temperature_ol"],
#           df_table_list_positions=t_ol[0],
#           df_ol_table=t_ol[1])
# print('')

# unload_ol(
#     template_path=SETTINGS["file_dir_temperature_template"],
#     template=SETTINGS["template_temperature"],
#     save_path=SETTINGS["dir_ol_save_directory"],
#     save_name=SETTINGS["file_name_temperature_ol"],
#     data_func=getTemptureForOL
# )


# def unload_ol_0(template_path, template_name, save_path, save_name, data_func):
#     temp_res = data_func()


# Заполнение таблиц построчно (без шаблона, df изначально соответсвует передаваемой таблице)

def write_table(table, df, not_centr_column=[], row_shift=0):
    # values = df.values
    count = table._column_count
    for row_idx in range(1, df.shape[0] + row_shift):
        table.add_row()
    clls = table._cells
    for row_idx in range(0, df.shape[0]):
        for column_idx in range(count):
            # clls[column_idx + row_idx * count].text = str(values[row_idx, column_idx])
            cell = clls[column_idx + (row_idx + row_shift) * count]
            cell.text = str(df.iloc[row_idx, column_idx])
            # Изменения стиля, шрифта и размера
            # cell_style(cell, 10)
            paragra_ph = cell.paragraphs[0]
            # paragra_ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            font = paragra_ph.runs[0].font
            font.name = 'Arial'
            font.size = Pt(10)

            # Выравнивание всех столбцов, кроме not_centr_column
            if column_idx not in not_centr_column:
                # paragra_ph = cell.paragraphs[0]
                paragra_ph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


# Выгрузка спецификации

def unloading_doc(template_path, save_path, save_name, data_func, not_centr_column=[], row_shift=0):
    full_save_path = os.path.join(save_path, save_name + '.docx')
    document = Document(template_path)
    table = document.tables[-1]
    df = data_func()
    write_table(table, df, not_centr_column, row_shift)
    document.save(full_save_path)
    # print(f'Документ {save_name}.docx выгружен')


# def

# Работает!
# unloading_doc(
#     template_path=SETTINGS['file_dir_specification_template'],
#     save_path=SETTINGS['dir_specifaication_save_directory'],
#     save_name=SETTINGS['file_name_specification'],
#     data_func=get_spec,
#     not_centr_column=[1]
# )

# Работает!
# unloading_doc(
#     template_path=SETTINGS['file_dir_io_template'],
#     save_path=SETTINGS['dir_io_save_directory'],
#     save_name=SETTINGS['file_name_io'],
#     data_func=get_io,
#     row_shift=2
# )

# Работает!
# unloading_doc(
#     template_path=SETTINGS['file_dir_tsp_template'],
#     save_path=SETTINGS['dir_tsp_save_directory'],
#     save_name=SETTINGS['file_name_tsp'],
#     data_func=get_tsp,
#     row_shift=1
# )
# unload_ol(template_path, template, save_path, save_name, data_func):

def unloading_kj(template_path, save_path, save_name, data_func):
    full_save_path = os.path.join(save_path, save_name + '.docx')
    document = Document(template_path)
    table_kj_list = document.tables[-1]
    table_kj_desc = document.tables[-2]
    df_list, df_desc = data_func()
    write_table(table_kj_list, df_list, row_shift=1)
    write_table(table_kj_desc, df_desc, row_shift=1)
    document.save(full_save_path)


# unloading_kj(
#     template_path=SETTINGS['file_dir_kj_template'],
#     save_path=SETTINGS['dir_kj_save_directory'],
#     save_name=SETTINGS['file_name_kj'],
#     data_func=get_kj
# )
